import gradio as gr
from llama_index.core import (
    Document as LlamaDocument,
    VectorStoreIndex,
)
from llama_index.core.tools import QueryEngineTool, ToolMetadata
from llama_index.llms.openai import OpenAI
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.core.agent.react.base import ReActAgent
from llama_index.core.node_parser import SentenceSplitter
from dotenv import load_dotenv
import os
import boto3
import mammoth
from docx import Document as DocxDocument
import pickle
from concurrent.futures import ThreadPoolExecutor, as_completed
from doctr.models import ocr_predictor, from_hub
from doctr.io import DocumentFile
import pymupdf4llm

from file_loader import FileLoader

model = from_hub("Felix92/doctr-torch-parseq-multilingual-v1")
predictor = ocr_predictor(
    det_arch="fast_base",
    reco_arch=model,
    pretrained=True,
    assume_straight_pages=True,
    detect_orientation=False,
)


def get_pdf_text(pdf_path: str):
    # docs = DocumentFile.from_pdf(pdf_path)
    # result = predictor(docs)
    return pymupdf4llm.to_markdown(pdf_path)
    # return str(result.render())


# Suppress warnings
import warnings

warnings.filterwarnings("ignore")

load_dotenv()
s3 = boto3.client("s3")

# Set up llama_index settings
embed_model = OpenAIEmbedding(model="text-embedding-3-large")
llm = OpenAI(temperature=0.9, model="gpt-4o")


def download_resumes_from_s3(bucket_name, local_dir, prefix=""):
    if not os.path.exists(local_dir):
        os.makedirs(local_dir)
    paginator = s3.get_paginator("list_objects_v2")
    pages = paginator.paginate(Bucket=bucket_name, Prefix=prefix)

    for page in pages:
        for obj in page.get("Contents", []):
            key = obj["Key"]
            if key.lower().endswith((".pdf", ".docx", ".doc", ".txt")):
                file_path = os.path.join(local_dir, key)
                os.makedirs(os.path.dirname(file_path), exist_ok=True)
                if os.path.exists(file_path):
                    print(f"Skipped: {key}")
                else:
                    s3.download_file(bucket_name, key, file_path)
                    print(f"Downloaded: {key}")


def load_document(file_path):
    if file_path.lower().endswith(".pdf"):
        return load_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return load_docx(file_path)
    elif file_path.lower().endswith(".doc"):
        return load_doc(file_path)
    elif file_path.lower().endswith(".txt"):
        return load_txt(file_path)
    else:
        print(f"Unsupported file type: {file_path}")
        return ""


def load_pdf(file_path):
    try:
        return get_pdf_text(file_path)
    except Exception as e:
        print(f"Error processing PDF file {file_path}: {e}")
        return ""


def load_docx(file_path):
    try:
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error processing .docx file {file_path}: {e}")
        return ""


def load_doc(file_path):
    try:
        with open(file_path, "rb") as doc_file:
            result = mammoth.extract_raw_text(doc_file)
        return result.value
    except Exception as e:
        print(f"Error processing .doc file {file_path}: {e}")
        return ""


def load_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    except Exception as e:
        print(f"Error processing .txt file {file_path}: {e}")
        return ""


def split_text(text, max_tokens=1024, chunk_overlap=64):
    splitter = SentenceSplitter(chunk_size=max_tokens, chunk_overlap=chunk_overlap)
    return splitter.split_text(text)


def process_file(file_path):
    documents = []
    print(file_path)
    loader = FileLoader(verbose=True)
    try:
        if file_path.lower().endswith(".pdf"):
            documents.extend(loader.load_pdf(file_path))
        elif file_path.lower().endswith(".docx"):
            documents.extend(loader.load_docx(file_path))
        elif file_path.lower().endswith(".doc"):
            text = load_doc(file_path)
            if text:
                text_chunks = split_text(text)
                for chunk in text_chunks:
                    documents.append(
                        LlamaDocument(text=chunk, metadata={"source": file_path})
                    )
        elif file_path.lower().endswith(".txt"):
            text = load_txt(file_path)
            text_chunks = split_text(text)
            for chunk in text_chunks:
                documents.append(
                    LlamaDocument(text=chunk, metadata={"source": file_path})
                )
    except Exception as e:
        print(f"Error processing file {file_path}: {e}")
    return documents


def preprocess_and_store_documents():
    # download_resumes_from_s3(os.environ["S3_BUCKET_NAME"], "resumes")

    all_documents = []
    with ThreadPoolExecutor(max_workers=1) as executor:
        futures = [
            executor.submit(process_file, os.path.join(root, file))
            for root, _, files in os.walk("resumes")
            for file in files
        ]
        for future in as_completed(futures):
            result = future.result()
            if result:
                all_documents.extend(result)

    texts = [d.text for d in all_documents if d.text.strip()]

    return all_documents, texts


def create_advanced_rag_system(documents):
    # Create a VectorStoreIndex
    index = VectorStoreIndex.from_documents(documents)

    # Create a base query engine
    base_query_engine = index.as_query_engine(
        similarity_top_k=100,
    )

    # Create query engine tools
    tools = [
        QueryEngineTool(
            query_engine=base_query_engine,
            metadata=ToolMetadata(
                name="vector_index",
                description="A vector database containing all the embeddings derived from each candidate's resume",
            ),
        )
    ]

    # Create the LLM
    llm = OpenAI(temperature=0.9, model="gpt-4o")

    # Create the RAG agent
    rag_agent = ReActAgent.from_tools(
        tools, llm=llm, verbose=True, react_chat_kwargs={"max_iterations": 5}
    )

    return rag_agent


def load_or_create_documents():
    if os.path.exists("documents.pkl"):
        with open("documents.pkl", "rb") as f:
            try:
                documents = pickle.load(f)
                if not documents:
                    raise EOFError
                return documents
            except EOFError:
                return None
    return None


def load_or_create_advanced_rag_system():
    documents = load_or_create_documents()
    if not documents:
        documents, _ = preprocess_and_store_documents()
        with open("documents.pkl", "wb") as f:
            pickle.dump(documents, f)

    # Ensure the documents have unique IDs
    for i, doc in enumerate(documents):
        doc.metadata["id"] = f"doc_{i}"

    # Create the advanced RAG system
    advanced_rag = create_advanced_rag_system(documents)

    return advanced_rag


# Initialize the advanced RAG system once at the start
advanced_rag = load_or_create_advanced_rag_system()


def add_text(history, text):
    if not text:
        raise gr.Error("Enter text")
    history.append((text, ""))
    return history, text


def generate_response(history, query):
    global advanced_rag

    chat_history_str = []
    for entry in history:
        if isinstance(entry, tuple) and len(entry) == 2:
            chat_history_str.append(entry)
        elif isinstance(entry, list) and len(entry) == 2:
            chat_history_str.append(tuple(entry))

    # Use the RAG agent
    rag_response = advanced_rag.chat(query)

    # Get the response
    response = f"RAG Agent: {rag_response}"

    history[-1] = (query, response)
    return history, ""


# Gradio application setup
with gr.Blocks(
    css="""
#chatbot {
  height: 600px;
  overflow-y: auto;
}
"""
) as demo:
    with gr.Column():
        chatbot = gr.Chatbot(value=[], elem_id="chatbot")
        txt = gr.Textbox(show_label=False, placeholder="Enter text and press enter")
        submit_btn = gr.Button("Submit")

    def submit_interaction(chatbot_value, txt_value):
        chatbot_value, query = add_text(chatbot_value, txt_value)
        chatbot_value, token_info = generate_response(chatbot_value, query)
        return chatbot_value, "", token_info

    submit_btn.click(
        fn=submit_interaction,
        inputs=[chatbot, txt],
        outputs=[chatbot, txt, gr.Textbox()],
    )

    txt.submit(
        fn=submit_interaction,
        inputs=[chatbot, txt],
        outputs=[chatbot, txt, gr.Textbox()],
    )

demo.queue()
if __name__ == "__main__":
    demo.launch(share=True)
